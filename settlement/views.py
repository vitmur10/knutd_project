from datetime import datetime
from django.http import Http404, FileResponse, HttpResponseRedirect
from django.shortcuts import render, redirect
from docx import Document
from docx.shared import Pt, Cm
from django.urls import reverse
from .forms import StudentForm

from .models import Hostel, Student

current_date = datetime.now()
formatted_month = current_date.strftime("%m")


def fill_order_document(template_path, output_path, data_to_fill, size, ):
    doc = Document(template_path)

    # Заповнення даними текстових параграфів
    for paragraph in doc.paragraphs:
        for key, value in data_to_fill.items():
            if key in paragraph.text:
                new_text = paragraph.text.replace(key, value)
                paragraph.text = new_text  # Замінюємо текст параграфу без використання run
                if paragraph.runs:
                    font = paragraph.runs[0].font
                    font.name = 'Times New Roman'
                    font.size = Pt(size)

    # Заповнення даними комірок таблиць

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data_to_fill.items():
                    if key in cell.text:
                        new_text = cell.text.replace(key, value)
                        cell.text = new_text  # Замінюємо текст комірки таблиці без використання run
                        if cell.paragraphs and cell.paragraphs[0].runs:
                            font = cell.paragraphs[0].runs[0].font
                            font.name = 'Times New Roman'
                            font.size = Pt(size)
                            # Виправлення відступів та розмірів комірки
                            cell.paragraphs[0].paragraph_format.left_indent = Cm(0.3)

    doc.save(output_path)


def add_student(request):
    if request.method == 'POST':
        form = StudentForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('settlement:four2')  # Перенаправити на сторінку зі списком студентів
    else:
        form = StudentForm()

    return render(request, 'swttlement/form.html', {'form': form})


def add_studentt(request):
    learning_from = "2021"  # Початок навчання
    training_to = "2025"  # Закінчення
    order_data_to_fill = {
        "м. Київ	  						«____»_______________20____р.": f"""м. Київ	  						"{current_date.day} {formatted_month} {current_date.year}" р.""",
        "виданий _____________________________________________________________________,": f"виданий                                            {request.POST['PIP']}",
        "який (яка) на вчається у __________________на денній формі навчання,": "який (яка) навчається у КНУТД на денній формі навчання,",
        "факультету _____________________________________,   _____курсу,    група___________,": f"{request.POST['faculty']},   {request.POST['course']} курсу,    група {request.POST['group']},",
        " № ____ по вул. _______________, буд.": f" № {request.POST['hostel']} по вул. {request.POST['street']}, буд. house № room, кімната, блок, № bloc",
        "№ ___, кімната, блок, № _________ .": "",
        "Ордер дійсний протягом ___________ навчального року до ___.______________ 20_____р.": f"Ордер дійсний протягом {learning_from}/{training_to} навчального року до 30.06.{training_to} р.",
        "_______________________________________________________,": f"                                                   {request.POST['IPI']},",
        "який (яка) навчається у _______________ на денній формі навчання,": "який (яка) навчається у КНУТД на денній формі навчання,",
        "факультету __________, __________курсу, група_______________": f"{request.POST['faculty']}, {request.POST['course']} курсу, група {request.POST['course']}",
        "на право займати ліжко-місце в гуртожитку №_____, кімнаті №___,": f"на право займати ліжко-місце в гуртожитку №{request.POST['hostel']}, кімнаті №{'room'},",
        "Термін проживання до __________________20_____ року": f"Термін проживання до 30.06.{training_to} року",
    }
    """            contract_data_to_fill = {
        " на	/	": f" на {[learning_from]}/{[training_to]}",
        "м. Київ	«	»	20	р.": f"м. Київ	«{current_date.day}» {formatted_month} {current_date.year} р.",
        "***********************************************************": f"{[IPI]}",
        "здобувача освіти факультету/інституту	група	курс	(далі — Наймач)": f"здобувача освіти факультету/інституту {[faculty]} група {[group]} курс {[course]} (далі — Наймач)",
        " в кімнаті №	у": f" в кімнаті №{room} у",
        "гуртожитку КНУТД №	, яке розташоване за адресою:	та зобов’язується провести ": f"гуртожитку КНУТД № {[hostel]}, яке розташоване за адресою: вул.{street} {house} та зобов’язується провести ",
        "Сплачувати за проживання. Оплата здійснюється авансовано за навчальний рік в сумі	гривень, без ПДВ.": f"Сплачувати за проживання. Оплата здійснюється авансовано за навчальний рік в сумі 10900 гривень, без ПДВ.",
        "Договір набуває чинності з   «     »	20           р. і діє до «      »	20 р": f"Договір набуває чинності з   «{current_date.day}» {formatted_month}. {current_date.year} р. і діє до «30» 06. {current_date.year + 1} р.",
        f"проживання в гуртожитку №  	": f"проживання в гуртожитку №  {[hostel]}",
        "Паспорт серія	№  	": f"Паспорт серія {[passport_series]} № {[passport_number]} ",
        "виданий «	»	20	р.": f"виданий «10»09 2020р. ",
        "ідентифікаційний код  	": f"ідентифікаційний код {[identification_code]}",
        "index, obl": f"38541 Хмельницька",
        "city": f"Старокостнятинів",
        " вулиця будинок квартира": f"вул. Лейпцезька 16 708",
        "?????????????????????????????????????????????????????????????????????????????????????????????????????????????????????????????": f"{[IPIF]}",
        "паспорт серія батьки	№ батьки	виданий  батьки	": f"паспорт серія {[passport_series_f]} № {[passport_number_f]} виданий {[passport_issued_f]}",
    }"""
    fill_order_document('ORDER.docx', f"заповнений_ордер{request.POST['IPI']}.docx", order_data_to_fill, 12)
    # fill_order_document(contract_template_path, contract_output_path, contract_data_to_fill, 8)
    filled_order_path = fill_order_document('ORDER.docx', f"заповнений_ордер{request.POST['IPI']}.docx",
                                            order_data_to_fill, 12)

    # Повертаємо файл для завантаження
    with open(filled_order_path, 'rb') as file:
        response = FileResponse(file,
                                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={filled_order_path.split("/")[-1]}'
    Student.objects.create(
        IPI=request.POST['PIP'],
        date_of_birth=request.POST['birth_date'],
        parents=request.POST['gender'],
        group=request.POST['group'],
        street=request.POST['street'],
        house=request.POST['house'],
        room=request.POST['room'],
        hostel=request.POST['dormitory'],
        faculty=request.POST['faculty'],
        course=request.POST['course'],
        bloc=request.POST['identification_index'],
        learning_from=request.POST['start_of_study'],
        training_to=request.POST['completion_of_study'],
        passport_series=request.POST['passport_series'],
        passport_number=request.POST['passport_number'],
        passport_issued=request.POST['passport_issued'],
        identification_code=request.POST['identification_index']
    )
    return render(request, 'swttlement/seven_hostel/floor_1.html')


def forms(request):
    return render(request, 'swttlement/form.html')


def dormitories(request):
    try:
        h = Hostel.objects.all()
    except:
        raise Http404('Сталася помилка(')
    return render(request, 'swttlement/dormitories.html', {'hostel': h})


class Hostel_seven:
    def one(request):
        return render(request, 'swttlement/seven_hostel/floor_1.html')

    def two(request):
        return render(request, 'swttlement/seven_hostel/floor_2.html')

    def three(request):
        return render(request, 'swttlement/seven_hostel/floor_3.html')

    def four(request):
        return render(request, 'swttlement/seven_hostel/floor_4.html')

    def five(request):
        return render(request, 'swttlement/seven_hostel/floor_5.html')

    def six(request):
        return render(request, 'swttlement/seven_hostel/floor_6.html')

    def seven(request):
        return render(request, 'swttlement/seven_hostel/floor_7.html')
