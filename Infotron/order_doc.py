from docx import Document
from datetime import datetime

from docx.shared import Pt, Cm

IPI = "Петров Іван Олександрович"  # ПІП
IPIF = "Галина Петрівна Олексіївна"  # ПіП Батьків
group = "БІТ2-21"  #Група
street = "Лейпцезька"  # Вулиця
house = "16"  # Дім
room = "708"  # Кімната
hostel = "7"  # Гуртожиток
faculty = "МКТ"  # Факультет
course = "3"  # Курс
bloc = "708"  # Блок
learning_from = "2021"  # Початок навчання
training_to = "2025"  # Закінчення
passport_series = '4545'  # Паспорт серія
passport_number = '45645653543'  # Паспорт номер
passport_issued = '76765'  # Паспорт виданий
identification_code = '687675736-6455'  # ІндиФІКАЦІЙНИЙ КОД
passport_series_f = '45645'  # Паспорт батьків серія
passport_number_f = '142373'  # Номер
passport_issued_f = '7878373'  #Виданий


def fill_order_document(template_path, output_path, data, size):
    doc = Document(template_path)

    # Заповнення даними текстових параграфів
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                new_text = paragraph.text.replace(key, value)
                paragraph.text = new_text  # Замінюємо текст параграфу без використання run
                if paragraph.runs:
                    font = paragraph.runs[0].font
                    font.name = 'Times New Roman'
                    font.size = Pt(size)

    # Заповнення даними комірок таблиць

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if key in cell.text:
                        new_text = cell.text.replace(key, value)
                        cell.text = new_text  # Замінюємо текст комірки таблиці без використання run
                        if cell.paragraphs and cell.paragraphs[0].runs:
                            font = cell.paragraphs[0].runs[0].font
                            font.name = 'Times New Roman'
                            font.size = Pt(size)
                            # Виправлення відступів та розмірів комірки
                            cell.paragraphs[0].paragraph_format.left_indent = Cm(0.3)

    doc.save(output_path)


current_date = datetime.now()
formatted_month = current_date.strftime("%m")
# Приклад використання
order_data_to_fill = {
    "м. Київ	  						«____»_______________20____р.": f"""м. Київ	  						"{current_date.day} {formatted_month} {current_date.year}" р.""",
    "виданий _____________________________________________________________________,": f"виданий                                            {IPI}",
    "який (яка) навчається у __________________на денній формі навчання,": "який (яка) навчається у КНУТД на денній формі навчання,",
    "факультету _____________________________________,   _____курсу,    група___________,": f"{faculty},   {course} курсу,    група {group},",
    " № ____ по вул. _______________, буд.": f" № {hostel} по вул. {street}, буд. {house} №{room}, кімната, блок, № {bloc}",
    "№ ___, кімната, блок, № _________ .": "",
    "Ордер дійсний протягом ___________ навчального року до ___.______________ 20_____р.": f"Ордер дійсний протягом {learning_from}/{training_to} навчального року до 30.06.{training_to} р.",
    "_______________________________________________________,": f"                                                   {IPI},",
    "який (яка) навчається у _______________ на денній формі навчання,": "який (яка) навчається у КНУТД на денній формі навчання,",
    "факультету __________, __________курсу, група_______________": f"{faculty}, {course} курсу, група {course}",
    "на право займати ліжко-місце в гуртожитку №_____, кімнаті №___,": f"на право займати ліжко-місце в гуртожитку №{hostel}, кімнаті №{room},",
    "Термін проживання до __________________20_____ року": f"Термін проживання до 30.06.{training_to} року",
}
contract_data_to_fill = {
    " на	/	": f" на {learning_from}/{training_to}",
    "м. Київ	«	»	20	р.": f"м. Київ	«{current_date.day}» {formatted_month} {current_date.year} р.",
    "***********************************************************": f"{IPI}",
    "здобувача освіти факультету/інституту	група	курс	(далі — Наймач)": f"здобувача освіти факультету/інституту {faculty} група {group} курс {course} (далі — Наймач)",
    " в кімнаті №	у": f" в кімнаті №{room} у",
    "гуртожитку КНУТД №	, яке розташоване за адресою:	та зобов’язується провести ": f"гуртожитку КНУТД № {hostel}, яке розташоване за адресою: вул.{street} {house} та зобов’язується провести ",
    "Сплачувати за проживання. Оплата здійснюється авансовано за навчальний рік в сумі	гривень, без ПДВ.": f"Сплачувати за проживання. Оплата здійснюється авансовано за навчальний рік в сумі 10900 гривень, без ПДВ.",
    "Договір набуває чинності з   «     »	20           р. і діє до «      »	20 р": f"Договір набуває чинності з   «{current_date.day}» {formatted_month}. {current_date.year} р. і діє до «30» 06. {current_date.year + 1} р.",
    f"проживання в гуртожитку №  	": f"проживання в гуртожитку №  {hostel}",
    "Паспорт серія	№  	": f"Паспорт серія {passport_series} № {passport_number} ",
    "виданий «	»	20	р.": f"виданий «10»09 2020р. ",
    "ідентифікаційний код  	": f"ідентифікаційний код {identification_code}",
    "index, obl": f"38541 Хмельницька",
    "city": f"Старокостнятинів",
    " вулиця будинок квартира": f"вул. Лейпцезька 16 708",
    "?????????????????????????????????????????????????????????????????????????????????????????????????????????????????????????????": f"{IPIF}",
    "паспорт серія батьки	№ батьки	виданий  батьки	": f"паспорт серія {passport_series_f} № {passport_number_f} виданий {passport_issued_f}",
}
order_template_path = "ОРДЕР студ наказ114-20 2021.docx"
order_output_path = "заповнений_ордер.docx"

fill_order_document(order_template_path, order_output_path, order_data_to_fill, 12)

contract_template_path = "Типовий_договір_найму_жилого_приміщення_На_Бугас_В_В_.docx"
contract_output_path = "заповнений_договір.docx"

fill_order_document(contract_template_path, contract_output_path, contract_data_to_fill, 8)
